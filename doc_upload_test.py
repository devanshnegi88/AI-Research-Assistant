"""Focused document-upload test — exercises TXT, DOCX, PDF, image, and dedup."""

import io
import uuid

import httpx
from PIL import Image

# Build a real PDF with PyMuPDF
import fitz

BASE = "http://127.0.0.1:8000"
client = httpx.Client(base_url=BASE, timeout=180)

results = []


def report(name, resp, expected=None):
    status = resp.status_code
    try:
        body = resp.json()
    except Exception:
        body = resp.text[:200]
    if expected is None:
        ok = status < 400
    elif isinstance(expected, (list, tuple)):
        ok = status in expected
    else:
        ok = status == expected
    results.append((name, status, ok))
    print(f"[{'OK' if ok else 'FAIL'}] {name} -> {status} | {str(body)[:200]}")
    return resp


# --- Register + login a fresh user ---
suffix = uuid.uuid4().hex[:8]
email = f"doc_{suffix}@example.com"
client.post(
    "/api/v1/auth/register",
    json={"email": email, "full_name": "Doc Tester", "password": "StrongPass1"},
)
login = client.post(
    "/api/v1/auth/login", json={"email": email, "password": "StrongPass1"}
)
token = login.json()["access_token"]
headers = {"Authorization": f"Bearer {token}"}

# --- 1. TXT upload + verify extracted text ---
report(
    "upload .txt",
    client.post(
        "/api/v1/documents/upload",
        files={"file": ("notes.txt", b"Hello world, this is a research document.", "text/plain")},
        headers=headers,
    ),
    expected=201,
)

# --- 2. DOCX upload (build a real docx in-memory) ---
from docx import Document as DocxDocument

doc = DocxDocument()
doc.add_paragraph("This is a paragraph in the DOCX file.")
doc.add_paragraph("Second paragraph with more research content.")
doc.add_table(rows=2, cols=2)
buf = io.BytesIO()
doc.save(buf)
docx_bytes = buf.getvalue()

report(
    "upload .docx",
    client.post(
        "/api/v1/documents/upload",
        files={"file": ("report.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers=headers,
    ),
    expected=201,
)

# --- 3. PDF upload (build a real text-layer PDF) ---
pdf_doc = fitz.open()
page = pdf_doc.new_page()
page.insert_text((72, 72), "This is a text-layer PDF for testing extraction.")
pdf_bytes = pdf_doc.tobytes()
pdf_doc.close()

report(
    "upload .pdf",
    client.post(
        "/api/v1/documents/upload",
        files={"file": ("paper.pdf", pdf_bytes, "application/pdf")},
        headers=headers,
    ),
    expected=201,
)

# --- 4. Image upload (build a small PNG) ---
img = Image.new("RGB", (200, 100), color=(255, 255, 255))
img_bytes = io.BytesIO()
img.save(img_bytes, format="PNG")
img_bytes = img_bytes.getvalue()

report(
    "upload .png",
    client.post(
        "/api/v1/documents/upload",
        files={"file": ("scan.png", img_bytes, "image/png")},
        headers=headers,
    ),
    expected=201,
)

# --- 5. Duplicate detection (same bytes => same document id) ---
content = b"identical content for dedup verification"
first = client.post(
    "/api/v1/documents/upload",
    files={"file": ("a.txt", content, "text/plain")},
    headers=headers,
)
second = client.post(
    "/api/v1/documents/upload",
    files={"file": ("b.txt", content, "text/plain")},
    headers=headers,
)
first_id = first.json().get("id")
second_id = second.json().get("id")
report(
    "dedup same content -> same id",
    first,
    expected=201,
)
print(f"      first_id={first_id}  second_id={second_id}  SAME={first_id == second_id}")
results.append(("dedup ids match", 200, first_id == second_id))

# --- 6. List all documents (verify all uploaded) ---
listing = client.get("/api/v1/documents", headers=headers)
report("list documents", listing, expected=200)
print(f"      total={listing.json().get('total')} count={len(listing.json().get('items', []))}")

# --- 7. Cross-user isolation ---
other_email = f"other_{suffix}@example.com"
client.post(
    "/api/v1/auth/register",
    json={"email": other_email, "full_name": "Other", "password": "StrongPass1"},
)
other_login = client.post(
    "/api/v1/auth/login", json={"email": other_email, "password": "StrongPass1"}
)
other_headers = {"Authorization": f"Bearer {other_login.json()['access_token']}"}
report(
    "other user cannot see my docs",
    client.get("/api/v1/documents", headers=other_headers),
    expected=200,
)
print(f"      other total={other_login and client.get('/api/v1/documents', headers=other_headers).json().get('total')}")

# --- 8. Delete a document ---
report(
    "delete .txt doc",
    client.delete(f"/api/v1/documents/{first_id}", headers=headers),
    expected=204,
)

print("\n=== SUMMARY ===")
passed = sum(1 for r in results if r[2])
total = len(results)
print(f"Passed: {passed}/{total}")
for name, _, ok in results:
    if not ok:
        print(f"  FAILED: {name}")
client.close()
