"""DOCX text + metadata extraction via python-docx."""

from __future__ import annotations

import io

from docx import Document as DocxDocument

from app.services.document.extractors.base import BaseExtractor, ExtractionResult


class DOCXExtractor(BaseExtractor):
    def extract(self, file_bytes: bytes) -> ExtractionResult:
        doc = DocxDocument(io.BytesIO(file_bytes))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        table_text: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_text.append(" | ".join(cells))

        full_text = "\n\n".join(paragraphs + table_text)
        metadata = self._build_metadata(doc)

        return ExtractionResult(text=full_text, metadata=metadata, used_ocr=False)

    def _build_metadata(self, doc: DocxDocument) -> dict:
        core_props = doc.core_properties
        return {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "title": core_props.title or None,
            "author": core_props.author or None,
            "created": core_props.created.isoformat() if core_props.created else None,
            "modified": core_props.modified.isoformat() if core_props.modified else None,
        }
